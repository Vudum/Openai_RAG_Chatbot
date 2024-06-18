import os
import time
import openai
import streamlit as st
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from requests_html import HTMLSession
from langchain.prompts import PromptTemplate
from langchain.memory import ConversationBufferMemory
from langchain.vectorstores import Chroma
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.llms import OpenAI
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from langchain.callbacks.manager import CallbackManager
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import RetrievalQA

# Set up environment variable and API key
openai_api_key = os.getenv("OPENAI_API_KEY")

# Ensure the API key is set
if openai_api_key is None:
    raise ValueError("OpenAI API key is not set. Please set it using the OPEN_API_KEY environment variable.")

# Use the API key for OpenAI
openai.api_key = openai_api_key

# Initialize directories
if not os.path.exists('uploadedFiles'):
    os.makedirs('uploadedFiles')

if not os.path.exists('vectorDB'):
    os.makedirs('vectorDB')

# Initialize session state variables
if 'template' not in st.session_state:
    st.session_state.template = """You are a knowledgeable chatbot, here to help with questions of the user. Your tone should be professional and informative.

    Context: {context}
    History: {history}

    User: {question}
    Chatbot:"""

if 'prompt' not in st.session_state:
    st.session_state.prompt = PromptTemplate(
        input_variables=["history", "context", "question"],
        template=st.session_state.template,
    )

if 'memory' not in st.session_state:
    st.session_state.memory = ConversationBufferMemory(
        memory_key="history",
        return_messages=True,
        input_key="question",
    )

if 'vectorstore' not in st.session_state:
    st.session_state.vectorstore = None

if 'llm' not in st.session_state:
    st.session_state.llm = OpenAI(api_key=openai_api_key,
                                  verbose=True,
                                  callback_manager=CallbackManager(
                                      [StreamingStdOutCallbackHandler()]),
                                  )

if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []

st.title("Chatbot - to talk to PDFs, Docs, Texts, and URLs")

uploaded_files = st.file_uploader("Choose files", type=["pdf", "docx", "txt"], accept_multiple_files=True)
url_input = st.text_input("Enter an internal URL")

for message in st.session_state.chat_history:
    with st.chat_message(message["role"]):
        st.markdown(message["message"])

all_documents = []

# Define a Document class to wrap content
class Document:
    def __init__(self, content, metadata=None):
        self.page_content = content
        self.metadata = metadata if metadata is not None else {}

# Function to handle document processing
def process_documents(uploaded_files):
    all_docs = []
    for uploaded_file in uploaded_files:
        file_path = os.path.join('uploadedFiles', uploaded_file.name)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        if uploaded_file.type == "application/pdf":
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
            all_docs.append(Document(text, {"source": uploaded_file.name}))
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = DocxDocument(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text
            all_docs.append(Document(text, {"source": uploaded_file.name}))
        elif uploaded_file.type == "text/plain":
            with open(file_path, "r") as f:
                text = f.read()
            all_docs.append(Document(text, {"source": uploaded_file.name}))

    return all_docs

# Process uploaded files
if uploaded_files:
    all_documents.extend(process_documents(uploaded_files))

# Process URL input
if url_input:
    session = HTMLSession()
    response = session.get(url_input)
    webpage_text = response.html.text.strip()  # Strip whitespace around text
    document = Document(webpage_text, {"source": url_input})  # Ensure webpage content is in expected format
    all_documents.append(document)

if all_documents:
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1500,
        chunk_overlap=200,
        length_function=len
    )

    all_splits = []
    for doc in all_documents:
        if hasattr(doc, "page_content"):
            splits = text_splitter.split_text(doc.page_content)
            for split in splits:
                all_splits.append(Document(split, doc.metadata))

    st.session_state.vectorstore = Chroma.from_documents(
        documents=all_splits,
        embedding=OpenAIEmbeddings(openai_api_key=openai_api_key)
    )

    st.session_state.vectorstore.persist()

    # Initialize retriever if it wasn't initialized before
    if 'vectorstore' in st.session_state and st.session_state.vectorstore:
        st.session_state.retriever = st.session_state.vectorstore.as_retriever()

    if 'qa_chain' not in st.session_state:
        st.session_state.qa_chain = RetrievalQA.from_chain_type(
            llm=st.session_state.llm,
            chain_type='stuff',
            retriever=st.session_state.retriever,
            verbose=True,
            chain_type_kwargs={
                "verbose": True,
                "prompt": st.session_state.prompt,
                "memory": st.session_state.memory,
            }
        )

user_question = st.text_input("Ask your question")

if user_question:
    user_message = {"role": "user", "message": user_question}
    st.session_state.chat_history.append(user_message)
    with st.chat_message("user"):
        st.markdown(user_question)

    with st.chat_message("assistant"):
        with st.spinner("Assistant is typing..."):
            response = st.session_state.qa_chain({"query": user_question, "context": ""})
        message_placeholder = st.empty()
        full_response = ""
        for chunk in response['result'].split():
            full_response += chunk + " "
            time.sleep(0.05)
            message_placeholder.markdown(full_response + "▌")
        message_placeholder.markdown(full_response)

    chatbot_message = {"role": "assistant", "message": response['result']}
    st.session_state.chat_history.append(chatbot_message)
